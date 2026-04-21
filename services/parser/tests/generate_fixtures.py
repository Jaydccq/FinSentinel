"""Generate test fixtures (PDF + DOCX) from the evaluation corpus.

This gives us deterministic, structure-ground-truth inputs for the
real extractors to validate against. Run once from the parser venv:

    cd services/parser
    .venv/bin/python tests/generate_fixtures.py

Writes to tests/fixtures/:
  - aapl-sample.pdf       (multi-heading, single-page, from chunks 1+3+5)
  - nvda-table.pdf        (multi-page with a 2D table, from chunks 9-11)
  - sample-memo.docx      (two headings + body text from chunk 20)

Fixtures are committed so CI can validate without the generator
running. Keep them small (< 50 KB each) to avoid repo bloat.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib import colors


REPO_ROOT = Path(__file__).resolve().parents[3]
CORPUS = REPO_ROOT / "services/evaluation-runner/datasets/corpus.json"
OUT = Path(__file__).parent / "fixtures"
OUT.mkdir(exist_ok=True)


def load_corpus() -> dict:
    return json.loads(CORPUS.read_text())


def _chunk_by_id(corpus: dict, chunk_id: str) -> str:
    for c in corpus["chunks"]:
        if c["chunk_id"] == chunk_id:
            return c["content"]
    raise KeyError(chunk_id)


def make_aapl_pdf(corpus: dict, out_path: Path) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER)
    flow = []
    # Title / PART / ITEM style to exercise the heading heuristic.
    flow.append(Paragraph("PART I", styles["Heading1"]))
    flow.append(Paragraph("ITEM 1. BUSINESS", styles["Heading2"]))
    flow.append(Paragraph(_chunk_by_id(corpus, "chunk-001"), styles["BodyText"]))
    flow.append(Spacer(1, 12))
    flow.append(Paragraph("ITEM 1A. RISK FACTORS", styles["Heading2"]))
    flow.append(Paragraph(_chunk_by_id(corpus, "chunk-003"), styles["BodyText"]))
    flow.append(Spacer(1, 12))
    flow.append(Paragraph("ITEM 7. LIQUIDITY AND CAPITAL RESOURCES", styles["Heading2"]))
    flow.append(Paragraph(_chunk_by_id(corpus, "chunk-005"), styles["BodyText"]))
    doc.build(flow)


def make_nvda_pdf_with_table(corpus: dict, out_path: Path) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER)
    flow = []
    flow.append(Paragraph("NVIDIA FY2025 RESULTS", styles["Heading1"]))
    flow.append(Paragraph(_chunk_by_id(corpus, "chunk-009"), styles["BodyText"]))
    flow.append(Spacer(1, 14))
    flow.append(Paragraph("SEGMENT REVENUE BREAKDOWN", styles["Heading2"]))
    data = [
        ["Segment", "FY2025 $B", "YoY %"],
        ["Data Center", "115.2", "+142%"],
        ["Gaming", "10.6", "+4%"],
        ["Professional Viz", "2.0", "+21%"],
        ["Automotive", "2.7", "+55%"],
    ]
    table = Table(data, colWidths=[180, 120, 120])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    flow.append(table)
    flow.append(PageBreak())
    flow.append(Paragraph("MARKET POSITION", styles["Heading2"]))
    flow.append(Paragraph(_chunk_by_id(corpus, "chunk-010"), styles["BodyText"]))
    doc.build(flow)


def make_memo_docx(corpus: dict, out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Microsoft AI Strategy Memo FY2025", level=1)
    doc.add_heading("Azure Growth", level=2)
    doc.add_paragraph(_chunk_by_id(corpus, "chunk-020"))
    doc.add_heading("Capital Intensity", level=2)
    doc.add_paragraph(_chunk_by_id(corpus, "chunk-021"))
    doc.add_heading("Summary Table", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "FY2025"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$245B"
    table.cell(2, 0).text = "Azure YoY"
    table.cell(2, 1).text = "+29%"
    doc.save(str(out_path))


def main() -> None:
    corpus = load_corpus()
    aapl_path = OUT / "aapl-sample.pdf"
    nvda_path = OUT / "nvda-table.pdf"
    docx_path = OUT / "sample-memo.docx"

    make_aapl_pdf(corpus, aapl_path)
    make_nvda_pdf_with_table(corpus, nvda_path)
    make_memo_docx(corpus, docx_path)

    for p in (aapl_path, nvda_path, docx_path):
        print(f"wrote {p.relative_to(REPO_ROOT)} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
