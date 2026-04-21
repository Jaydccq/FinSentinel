"""DOCX extractor backed by python-docx.

Walks document.paragraphs in order and maps each paragraph's style to a
Markdown heading level (Heading 1 → h1, Heading 2 → h2, etc.) or emits
plain text. Tables are rendered as Markdown pipes. Nested structures are
flattened — python-docx yields paragraphs and tables as separate elements
at the document level, which matches the Markdown contract the API
needs.
"""

from __future__ import annotations

import io
from typing import List, Tuple

from docx import Document

from .types import ExtractorResult, Heading


def extract_docx(data: bytes) -> ExtractorResult:
    doc = Document(io.BytesIO(data))

    md_parts: List[str] = []
    headings: List[Heading] = []
    table_count = 0

    # python-docx exposes .paragraphs and .tables, but they live in
    # different iteration orders. For faithful rendering we walk the
    # underlying body element order instead.
    for element in _iter_block_items(doc):
        if element is None:
            continue
        kind, payload = element
        if kind == "paragraph":
            level = _heading_level_from_style(payload.style.name if payload.style else "")
            text = (payload.text or "").strip()
            if not text:
                md_parts.append("")
                continue
            if level:
                headings.append(Heading(level=level, text=text, pageStart=None))
                md_parts.append(f"{'#' * level} {text}")
            else:
                md_parts.append(text)
        elif kind == "table":
            md_parts.append(_table_to_markdown(payload))
            table_count += 1

    page_count = _estimate_page_count(doc)

    return ExtractorResult(
        markdown="\n\n".join(p for p in md_parts if p is not None),
        headings=headings,
        page_count=page_count,
        table_count=table_count,
    )


def _iter_block_items(doc):
    """Yield (kind, element) for paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            # python-docx has a Paragraph constructor, but the paragraphs
            # are already lazily constructed via doc.paragraphs. Match by
            # element identity.
            for p in doc.paragraphs:
                if p._p is child:
                    yield ("paragraph", p)
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._tbl is child:
                    yield ("table", t)
                    break


def _heading_level_from_style(style_name: str) -> int | None:
    """Map 'Heading 1'..'Heading 6' to levels 1..6. 'Title' → 1."""
    if not style_name:
        return None
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading "):
        try:
            n = int(style_name.split(" ", 1)[1])
        except (ValueError, IndexError):
            return None
        if 1 <= n <= 6:
            return n
    return None


def _table_to_markdown(table) -> str:
    """Render a python-docx table as a GitHub-flavored Markdown table."""
    rows = [[cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
            for row in table.rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    header = "| " + " | ".join((rows[0] + [""] * width)[:width]) + " |"
    sep = "| " + " | ".join("---" for _ in range(width)) + " |"
    body = [
        "| " + " | ".join((r + [""] * width)[:width]) + " |"
        for r in rows[1:]
    ]
    return "\n".join([header, sep, *body])


def _estimate_page_count(doc) -> int:
    """python-docx has no native page count; estimate from pagination hints."""
    # Count explicit page breaks in runs; fall back to a rough 40-paragraphs
    # per page estimate. This is only advisory — the API does not depend on
    # exact page counts for retrieval.
    breaks = 1  # start with page 1
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for br in run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
            ):
                if br.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                ) == "page":
                    breaks += 1
    if breaks > 1:
        return breaks
    # Heuristic fallback
    para_count = len(doc.paragraphs)
    return max(1, para_count // 40 + (1 if para_count % 40 else 0))
